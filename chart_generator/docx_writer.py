import re
from typing import List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from .transposer import CHORD_REGEX, is_chord_line


# Regex to match a Nashville number chord token (e.g., 1, 4m, 5sus4, 6/5)
NASH_REGEX = re.compile(r"[1-7][^\s]*")


def is_nashville_line(line: str) -> bool:
    tokens = line.strip().split()
    return bool(tokens) and all(NASH_REGEX.fullmatch(tok) for tok in tokens)


MONO_FONT = "Courier New"


def export_chart_to_docx(lines: List[str], output_path: str):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = MONO_FONT
    font.size = Pt(10)

    for line in lines:
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = 1.15
        if is_chord_line(line) or is_nashville_line(line):
            last_idx = 0
            pattern = NASH_REGEX if is_nashville_line(line) else CHORD_REGEX
            for m in pattern.finditer(line):
                # lyric spacing before chord
                if m.start() > last_idx:
                    run = p.add_run(line[last_idx:m.start()])
                    run.font.name = MONO_FONT
                chord_run = p.add_run(m.group(0))
                chord_run.bold = True
                chord_run.font.name = MONO_FONT
                last_idx = m.end()
            # remainder
            if last_idx < len(line):
                r = p.add_run(line[last_idx:])
                r.font.name = MONO_FONT
        else:
            run = p.add_run(line)
            run.font.name = MONO_FONT
    doc.save(output_path)
